import os
import spacy
import streamlit as st
from pdfminer.high_level import extract_text as extract_text_from_pdf
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import subprocess
import sys

# Ensure the model is downloaded
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")  # Load model after installation

# Streamlit UI
st.title("📄 Multi-Resume Screening App")
st.write("Upload multiple resumes (PDF/DOCX) and enter a job description to rank them based on relevance.")

# User inputs job description
job_description = st.text_area("📌 Enter Job Description:")

# File uploader (Multiple files, limit set to 10)
uploaded_files = st.file_uploader("📂 Upload Resumes (Max 10)", type=["pdf", "docx"], accept_multiple_files=True)

# User selects the number of top resumes to display
top_n = st.number_input("🎯 Select Top N Resumes to Display", min_value=1, max_value=10, value=3, step=1)

if uploaded_files and job_description:
    if len(uploaded_files) > 10:
        st.warning("⚠️ You can only upload up to 10 resumes.")
    else:
        resume_scores = []

        for file in uploaded_files:
            # Extract text from PDF/DOCX
            if file.name.endswith(".pdf"):
                resume_text = extract_text_from_pdf(file).lower()
            else:
                doc = Document(file)
                resume_text = "\n".join([para.text for para in doc.paragraphs]).lower()

            # Extract keywords from resume & job description
            resume_keywords = {token.text for token in nlp(resume_text) if token.is_alpha and not token.is_stop}
            job_desc_keywords = {token.text for token in nlp(job_description) if token.is_alpha and not token.is_stop}

            # TF-IDF + Cosine Similarity
            vectorizer = TfidfVectorizer()
            tfidf_matrix = vectorizer.fit_transform([" ".join(job_desc_keywords), " ".join(resume_keywords)])
            score = round(cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0] * 100, 2)

            # Store scores
            resume_scores.append((file.name, score))

        # Sort resumes based on match score (Descending order)
        resume_scores.sort(key=lambda x: x[1], reverse=True)

        # Categorize resumes into High, Medium, and Low
        high_matches = [res for res in resume_scores if res[1] >= 70]
        medium_matches = [res for res in resume_scores if 40 <= res[1] < 70]
        low_matches = [res for res in resume_scores if res[1] < 40]

        
         # Display Top N resumes selected by the user
        st.subheader(f"🏆 **Top {top_n} Resume(s) Based on Match Score**")
        for i in range(min(top_n, len(resume_scores))):
            st.write(f"🥇 **{resume_scores[i][0]}**  ==>>   Match Score: {resume_scores[i][1]}%")

        
        # Display ranked resumes
        st.subheader("📊 Resume Ranking Results")

        # High Matches
        if high_matches:
           # st.success("✅ **High Match Resumes (≥ 70%)**")
            for name, score in high_matches:
                st.write(f"📌 **{name}**  → Match Score: {score}%")

        # Medium Matches
        if medium_matches:
           # st.warning("⚠️ **Medium Match Resumes (40% - 69%)**")
            for name, score in medium_matches:
                st.write(f"📌 **{name}**  → Match Score: {score}%")

        # Low Matches
        if low_matches:
            #st.error("❌ **Low Match Resumes (< 40%)**")
            for name, score in low_matches:
                st.write(f"📌 **{name}** → Match Score: {score}%")
        # Extract matched keywords
        matched_keywords = list(resume_keywords & job_desc_keywords)



# Add a section for missing (unmatched) keywords
st.subheader("❌ Keywords Missing in Each Resume")

for file, score in resume_scores:
    # Extract missing keywords (keywords in job description but not in resume)
    missing_keywords = list(job_desc_keywords - resume_keywords_dict[file])
    
    with st.expander(f"📄 {file} - Missing Keywords"):
        if missing_keywords:
            st.write(", ".join(missing_keywords))
        else:
            st.write("✅ No missing keywords! This resume covers all job description terms.")

        # # Checkbox to show matched keywords
        # if st.checkbox("🔍 Show Matched Keywords"):
        #     if matched_keywords:
        #         st.write(", ".join(matched_keywords))
        #     else:
        #         st.write("No matched keywords found.")




#import os
# import spacy
# import streamlit as st
# from pdfminer.high_level import extract_text as extract_text_from_pdf
# from docx import Document
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity

# # Load NLP model
# try:
#     nlp = spacy.load("en_core_web_sm")
# except OSError:
#     st.error("Spacy model 'en_core_web_sm' not found. Run `python -m spacy download en_core_web_sm` to install it.")
#     st.stop()

# # Streamlit UI
# st.title("📄 Multi-Resume Screening App")
# st.write("Upload multiple resumes (PDF/DOCX) and enter a job description to rank them based on relevance.")

# # User inputs job description
# job_description = st.text_area("📌 Enter Job Description:")

# # File uploader (Multiple files, limit set to 10)
# uploaded_files = st.file_uploader("📂 Upload Resumes (Max 10)", type=["pdf", "docx"], accept_multiple_files=True)

# if uploaded_files and job_description:
#     if len(uploaded_files) > 10:
#         st.warning("⚠️ You can only upload up to 10 resumes.")
#     else:
#         resume_scores = []

#         for file in uploaded_files:
#             # Extract text from PDF/DOCX
#             if file.name.endswith(".pdf"):
#                 resume_text = extract_text_from_pdf(file).lower()
#             else:
#                 doc = Document(file)
#                 resume_text = "\n".join([para.text for para in doc.paragraphs]).lower()

#             # Extract keywords from resume & job description
#             resume_keywords = {token.text for token in nlp(resume_text) if token.is_alpha and not token.is_stop}
#             job_desc_keywords = {token.text for token in nlp(job_description) if token.is_alpha and not token.is_stop}

#             # TF-IDF + Cosine Similarity
#             vectorizer = TfidfVectorizer()
#             tfidf_matrix = vectorizer.fit_transform([" ".join(job_desc_keywords), " ".join(resume_keywords)])
#             score = round(cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0] * 100, 2)

#             # Store scores
#             resume_scores.append((file.name, score))

#         # Sort resumes based on match score (Descending order)
#         resume_scores.sort(key=lambda x: x[1], reverse=True)

#         # Categorize resumes into High, Medium, and Low
#         high_matches = [res for res in resume_scores if res[1] >= 70]
#         medium_matches = [res for res in resume_scores if 40 <= res[1] < 70]
#         low_matches = [res for res in resume_scores if res[1] < 40]

#         # Display ranked resumes
#         st.subheader("📊 Resume Ranking Results")

#         # High Matches
#         if high_matches:
#             st.success("✅ **High Match Resumes (≥ 70%)**")
#             for name, score in high_matches:
#                 st.write(f"📌 **{name}** → Match Score: {score}%")

#         # Medium Matches
#         if medium_matches:
#             st.warning("⚠️ **Medium Match Resumes (40% - 69%)**")
#             for name, score in medium_matches:
#                 st.write(f"📌 **{name}** → Match Score: {score}%")

#         # Low Matches
#         if low_matches:
#             st.error("❌ **Low Match Resumes (< 40%)**")
#             for name, score in low_matches:
#                 st.write(f"📌 **{name}** → Match Score: {score}%")

